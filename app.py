import streamlit as st
import json
import os
from docx import Document
import requests
import base64

# --- GitHub repo details from environment variables ---
GITHUB_TOKEN = os.environ.get("GITHUB_TOKEN")
GITHUB_REPO = os.environ.get("GITHUB_REPO")
GITHUB_FILE = os.environ.get("GITHUB_FILE")

# GitHub API base URL
GITHUB_API = f"https://api.github.com/repos/{GITHUB_REPO}/contents/{GITHUB_FILE}"

# --- Function: Load file from GitHub ---
def load_error_db():
    headers = {"Authorization": f"token {GITHUB_TOKEN}"}
    res = requests.get(GITHUB_API, headers=headers)
    if res.status_code == 200:
        content = res.json()
        data = base64.b64decode(content["content"]).decode("utf-8")
        return json.loads(data), content["sha"]
    else:
        # If file doesn't exist, return empty dict
        return {}, None

# --- Function: Save file to GitHub ---
def save_error_db(error_db, file_sha=None):
    url = f"https://api.github.com/repos/{GITHUB_REPO}/contents/{GITHUB_FILE}"
    headers = {"Authorization": f"token {GITHUB_TOKEN}"}

    content = base64.b64encode(json.dumps(error_db, indent=2).encode()).decode()

    data = {
        "message": "Update error DB",
        "content": content
    }
    if file_sha:  # only include if updating an existing file
        data["sha"] = file_sha

    res = requests.put(url, headers=headers, json=data)

    if res.status_code in (200, 201):
        st.success("✅ Error DB saved successfully!")
    else:
        try:
            error_json = res.json()   # if valid JSON
            st.error(f"❌ Failed to save data: {error_json}")
        except ValueError:
            st.error(f"❌ Failed to save data: {res.status_code} - {res.text}")


# --- Load DB at startup ---
error_db, file_sha = load_error_db()

# --- Store uploaded files in JSON every run ---
LOCAL_JSON_FILE = "error_data_local.json"
if os.path.exists(LOCAL_JSON_FILE):
    with open(LOCAL_JSON_FILE, "r", encoding="utf-8") as f:
        local_db = json.load(f)
else:
    local_db = {}

# Merge GitHub DB with local DB
error_db.update(local_db)

# --- File uploader ---
uploaded_files = st.file_uploader("Upload Word Documents", type=["docx"], accept_multiple_files=True)

if uploaded_files:
    for uploaded_file in uploaded_files:
        doc = Document(uploaded_file)
        # Use first paragraph as title
        title = doc.paragraphs[0].text.strip() if doc.paragraphs else uploaded_file.name
        content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        error_db[title] = content

    # Save to GitHub
    save_error_db(error_db, file_sha)
    
    # Also save locally every run
    with open(LOCAL_JSON_FILE, "w", encoding="utf-8") as f:
        json.dump(error_db, f, ensure_ascii=False, indent=4)
    st.success("✅ Data stored locally and on GitHub!")

# --- Search ---
search_title = st.text_input("Enter the error title to search")

if search_title:
    if search_title in error_db:
        st.subheader(f"📄 {search_title}")
        # Format into steps
        steps = error_db[search_title].split("\n")
        for i, step in enumerate(steps, 1):
            if step.strip():
                st.markdown(f"**{i})** {step.strip()}")
    else:
        st.warning("No details found for this title.")
